#!/usr/bin/env python3
"""
创建测试文档的脚本
"""

from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

def create_test_word_document():
    """创建测试Word文档"""
    doc = Document()
    
    doc.add_heading('人工智能教学材料', 0)
    
    doc.add_heading('第一章：人工智能概述', level=1)
    
    doc.add_paragraph('人工智能（Artificial Intelligence，简称AI）是计算机科学的一个重要分支。')
    
    doc.add_paragraph('人工智能的核心目标是创造能够模拟人类智能行为的机器系统。')
    
    doc.add_heading('主要特征', level=2)
    doc.add_paragraph('1. 学习能力：通过数据和经验不断改进')
    doc.add_paragraph('2. 推理能力：基于逻辑和规则进行决策')
    doc.add_paragraph('3. 适应能力：根据环境变化调整行为')
    doc.add_paragraph('4. 创新能力：产生新的解决方案')
    
    doc.add_heading('应用领域', level=2)
    doc.add_paragraph('• 机器学习和深度学习')
    doc.add_paragraph('• 计算机视觉和图像识别')
    doc.add_paragraph('• 自然语言处理')
    doc.add_paragraph('• 机器人技术')
    doc.add_paragraph('• 智能决策系统')
    
    doc.add_heading('发展趋势', level=2)
    doc.add_paragraph('人工智能技术正在快速发展，预计在未来几年将在以下方面取得重大突破：')
    doc.add_paragraph('1. 更强的泛化能力')
    doc.add_paragraph('2. 更高的计算效率')
    doc.add_paragraph('3. 更好的人机交互')
    doc.add_paragraph('4. 更广泛的应用场景')
    
    doc.save('test_ai_document.docx')
    print("Word文档已创建：test_ai_document.docx")

def create_test_pdf_document():
    """创建测试PDF文档"""
    filename = "test_ai_document.pdf"
    c = canvas.Canvas(filename, pagesize=letter)
    width, height = letter
    
    # 添加内容
    c.setFont("Helvetica-Bold", 16)
    c.drawString(100, height - 80, "Artificial Intelligence Course Material")
    
    c.setFont("Helvetica-Bold", 14)
    c.drawString(100, height - 120, "Chapter 1: Introduction to AI")
    
    c.setFont("Helvetica", 12)
    y_position = height - 160
    
    content = [
        "Artificial Intelligence (AI) is a branch of computer science that aims to create",
        "intelligent machines that can think and act like humans. AI systems are designed",
        "to perform tasks that typically require human intelligence.",
        "",
        "Key Features of AI:",
        "1. Learning - The ability to improve performance through experience",
        "2. Reasoning - Using logic and rules to reach conclusions",
        "3. Perception - Understanding and interpreting sensory information",
        "4. Language Processing - Understanding and generating human language",
        "",
        "Applications of AI:",
        "• Machine Learning and Deep Learning",
        "• Computer Vision and Image Recognition",
        "• Natural Language Processing",
        "• Robotics and Automation",
        "• Expert Systems and Decision Support",
        "",
        "Future Developments:",
        "AI technology is rapidly advancing and is expected to achieve breakthroughs in:",
        "- Enhanced generalization capabilities",
        "- Improved computational efficiency", 
        "- Better human-machine interaction",
        "- Wider range of applications"
    ]
    
    for line in content:
        if line.startswith("Key Features") or line.startswith("Applications") or line.startswith("Future"):
            c.setFont("Helvetica-Bold", 12)
        else:
            c.setFont("Helvetica", 11)
        
        c.drawString(100, y_position, line)
        y_position -= 20
        
        if y_position < 100:
            c.showPage()
            y_position = height - 80
    
    c.save()
    print(f"PDF文档已创建：{filename}")

if __name__ == "__main__":
    try:
        create_test_word_document()
        create_test_pdf_document()
        print("测试文档创建完成！")
    except Exception as e:
        print(f"创建文档时出错：{e}")